import os
from docx import Document
from datetime import datetime
from io import BytesIO

MESES = {
    1: "janeiro", 2: "fevereiro", 3: "março", 4: "abril",
    5: "maio", 6: "junho", 7: "julho", 8: "agosto",
    9: "setembro", 10: "outubro", 11: "novembro", 12: "dezembro"
}


def normalizar(texto):
    if not texto:
        return ""
    return (
        str(texto)
        .lower()
        .replace(" ", "_")
        .replace("ã", "a")
        .replace("á", "a")
        .replace("é", "e")
        .replace("í", "i")
        .replace("ó", "o")
        .replace("ú", "u")
    )


def formatar_data(data_str):
    if not data_str:
        return ""
    try:
        return datetime.strptime(data_str, "%Y-%m-%d").strftime("%d/%m/%Y")
    except:
        return data_str


def mapear_student_para_word(dados):
    responsaveis = dados.get("responsaveisLegais", [])

    resp1 = responsaveis[0] if len(responsaveis) > 0 else {}
    resp2 = responsaveis[1] if len(responsaveis) > 1 else {}

    logradouro = dados.get("enderecoLogradouro") or ""
    numero = dados.get("enderecoNumero") or ""

    # 🔥 FIX ESTADO CIVIL (ESSENCIAL)
    MAP_ESTADO_CIVIL = {
        "casado": "casado",
        "uniao_estavel": "uniao",
        "separados": "separado",
        "divorciados": "divorciado",
        "viuvo": "viuvo",
        "outro": "outro"
    }

    return {
        # ===== DADOS PRINCIPAIS =====
        "nome_crianca": dados.get("nomeCompleto", ""),
        "data_nascimento": formatar_data(dados.get("dataNascimento")),
        "idade": dados.get("idade", ""),
        "naturalidade": dados.get("naturalidade", ""),
        "raca_cor": dados.get("racaCor", ""),
        "sexo": dados.get("sexo", ""),
        "rg_crianca": dados.get("rg", ""),
        "cpf_crianca": dados.get("cpf", ""),
        "nis": dados.get("nis", ""),
        "cras_referencia": dados.get("crasReferencia", ""),

        # ===== ENDEREÇO =====
        "endereco": f"{logradouro}, {numero}".strip(", "),
        "bairro": dados.get("enderecoBairro", ""),
        "cidade": dados.get("enderecoCidade", ""),
        "cep": dados.get("enderecoCep", ""),

        # ===== PAIS =====
        "nome_pai": dados.get("nomePai", ""),
        "nome_mae": dados.get("nomeMae", ""),

        # ===== RESPONSÁVEIS =====
        "responsavel_1_nome": resp1.get("nome", ""),
        "responsavel_1_rg": resp1.get("rg", ""),
        "responsavel_1_cpf": resp1.get("cpf", ""),
        "responsavel_1_celular": resp1.get("celular", ""),
        "responsavel_1_parentesco": resp1.get("parentesco", ""),

        "responsavel_2_nome": resp2.get("nome", ""),
        "responsavel_2_rg": resp2.get("rg", ""),
        "responsavel_2_cpf": resp2.get("cpf", ""),
        "responsavel_2_celular": resp2.get("celular", ""),
        "responsavel_2_parentesco": resp2.get("parentesco", ""),

        # ===== ESCOLAR =====
        "escola": dados.get("escola", ""),
        "serie": dados.get("serie", ""),
        "periodo_escolar": dados.get("periodo_escolar", ""),

        # ===== SAÚDE =====
        "ubs_referencia": dados.get("ubs_referencia", ""),
        "medicamento_continuo": dados.get("medicamento_continuo", ""),
        "alergia_qual": dados.get("alergia_qual", ""),
        "problema_saude_qual": dados.get("problema_saude_qual", ""),
        "restricao_alimentar_qual": dados.get("restricao_alimentar_qual", ""),
        "restricao_fisica_qual": dados.get("restricao_fisica_qual", ""),
        "deficiencia_qual": dados.get("deficiencia_qual", ""),

        # ===== CAMPOS CONTROLADOS =====
        "origem": normalizar(dados.get("origem")),
        "estado_civil": MAP_ESTADO_CIVIL.get(
            normalizar(dados.get("estado_civil")),
            normalizar(dados.get("estado_civil"))
        ),
        "tipo_domicilio": normalizar(dados.get("tipo_domicilio")),
        "vai": normalizar(dados.get("vai")),

        # ===== BOOLEANOS =====
        "matriculado": dados.get("matriculado"),
        "parou_escola": dados.get("parou_escola"),
        "problema_saude": dados.get("problema_saude"),
        "restricao_alimentar": dados.get("restricao_alimentar"),
        "restricao_fisica": dados.get("restricao_fisica"),
        "bronquite": dados.get("bronquite"),
        "falta_ar": dados.get("falta_ar"),
        "odontologico": dados.get("odontologico"),
        "deficiencia": dados.get("deficiencia"),
        "oftalmologico": dados.get("oftalmologico"),
        "usa_oculos": dados.get("usa_oculos"),
        "fica_sozinho": dados.get("fica_sozinho"),
        "outras_atividades": dados.get("outras_atividades"),
        "situacao_prioritaria": dados.get("situacao_prioritaria"),

        # ===== LISTAS =====
        "beneficios": [normalizar(b) for b in dados.get("beneficios", [])],
        "onde": [normalizar(x) for x in dados.get("onde", [])],
        "atividade": [normalizar(x) for x in dados.get("atividade", [])],
        "servicos": [normalizar(x) for x in dados.get("servicos", [])],
        "atendimentos": [normalizar(x) for x in dados.get("atendimentos", [])],

        # ===== INTERAÇÃO =====
        "interage_frequencia": normalizar(dados.get("interage_frequencia")),
        "interage_com": [normalizar(x) for x in dados.get("interage_com", [])],
    }


def preencher_documento(nome_arquivo, dados):
    if not os.path.exists(nome_arquivo):
        raise FileNotFoundError(f"Template não encontrado: {nome_arquivo}")

    doc = Document(nome_arquivo)

    hoje = datetime.now()
    dados.setdefault('dia', hoje.strftime('%d'))
    dados.setdefault('mes', MESES[hoje.month])
    dados.setdefault('ano', hoje.strftime('%Y'))

    def substituir(paragraph):
        texto = "".join(run.text for run in paragraph.runs)

        for chave, valor in dados.items():
            texto = texto.replace(f"{{{chave}}}", str(valor))

        if paragraph.runs:
            paragraph.runs[0].text = texto
            for run in paragraph.runs[1:]:
                run.text = ""

    for p in doc.paragraphs:
        substituir(p)

    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for p in celula.paragraphs:
                    substituir(p)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer